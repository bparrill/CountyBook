from   docx   import Document
import pandas as     pd

import math

# County list from https://www.census.gov/data/datasets/time-series/demo/popest/2020s-counties-total.html

df = pd.read_excel('co-est2021-pop.xlsx', header=None, skiprows=5, skipfooter=5)
df.columns =['county_st', 'pop', 'pop1', 'pop2']
df = df[['county_st', 'pop']]
df[['county', 'state']] = df.county_st.str.split(', ', expand=True)
# df['county'] = df.county.str[1:-7]
df['county'] = df['county'].apply(lambda x: x[1:-7] if x.find(' County') > -1 else x[1:])
df.sort_values(['state', 'county'], inplace=True)
df['county'] = chr(9633) + ' ' + df['county'] + '\n'
df_states = df.groupby("state")

# print(df)

document = Document('county_template.docx')
#document.add_heading('Document Title', 0)
# Unicode test:
# document.add_paragraph(f'Box: {chr(9633)}')

# For each state:
#    - New Page
#    - Page title (State name level 1)
#    - Flag?
#    - List of counties
#    - Total number of counties
#    - State facts?

for state in sorted(df.state.unique()):
    num_counties = len(df_states.get_group(state).county)
    print(f'{state} ({num_counties})')
    document.add_page_break()  
    document.add_paragraph(state, style='Heading')
    cols = min((num_counties // 35) + 1, 6)
    per_col = math.ceil(num_counties / cols)
    table = document.add_table(rows=1, cols=cols)
    row = table.row_cells(0)
    breaks = list(range(0, num_counties, per_col))
    if breaks[-1] != num_counties:
        breaks.append(num_counties)
    for col in range(len(breaks)-1):
        row[col].add_paragraph("".join(df_states.get_group(state).county.iloc[breaks[col]:breaks[col+1]]))
    document.add_paragraph(f'Counties in {state}: {num_counties}')

document.save('demo.docx')
